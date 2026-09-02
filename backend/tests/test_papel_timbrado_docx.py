"""
Testes de gerar_docx_timbrado (papel timbrado em .docx pra declarações
redigidas pelo usuário) — mesma identidade visual da proposta em PDF, mas
corpo vazio. Rode com:  cd backend && pytest
"""
import io

from docx import Document
from docx.oxml.ns import qn

from app.papel_timbrado_docx import gerar_docx_timbrado

_LOGO_PNG = ("data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4"
            "2mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=")


def _texto_header(doc: Document) -> str:
    header = doc.sections[0].header
    partes = [p.text for p in header.paragraphs]
    for tabela in header.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                partes.extend(p.text for p in celula.paragraphs)
    return "\n".join(partes)


def _texto_footer(doc: Document) -> str:
    return "\n".join(p.text for p in doc.sections[0].footer.paragraphs)


def _tem_imagem_no_header(doc: Document) -> bool:
    rels = doc.sections[0].header.part.rels
    return any(rel.reltype.endswith("/image") for rel in rels.values())


def _remetente(**over):
    base = {
        "nome": "Empresa Exemplo LTDA",
        "documento": "12345678000199",
        "endereco": {"logradouro": "Rua das Flores", "numero": "100", "bairro": "Centro",
                    "cidade": "Maringá", "uf": "PR", "cep": "87000-000"},
        "empresa": {"telefone": "(44) 99999-0000"},
        "logo_base64": None,
    }
    base.update(over)
    return base


def test_gera_docx_valido_com_nome_e_documento_no_cabecalho():
    conteudo = gerar_docx_timbrado(_remetente())
    doc = Document(io.BytesIO(conteudo))
    cabecalho = _texto_header(doc)
    assert "Empresa Exemplo LTDA" in cabecalho
    assert "12345678000199" in cabecalho


def test_rodape_junta_telefone_e_endereco():
    conteudo = gerar_docx_timbrado(_remetente())
    doc = Document(io.BytesIO(conteudo))
    rodape = _texto_footer(doc)
    assert "(44) 99999-0000" in rodape
    assert "Maringá" in rodape
    assert "PR" in rodape
    assert "87000-000" in rodape


def test_corpo_do_documento_fica_vazio_pra_usuario_redigir():
    conteudo = gerar_docx_timbrado(_remetente())
    doc = Document(io.BytesIO(conteudo))
    texto_corpo = "\n".join(p.text for p in doc.paragraphs).strip()
    assert texto_corpo == ""


def test_remetente_minimo_sem_endereco_nem_documento_nao_quebra():
    conteudo = gerar_docx_timbrado({"nome": "Fulano", "documento": None,
                                    "endereco": {}, "empresa": {}, "logo_base64": None})
    doc = Document(io.BytesIO(conteudo))
    assert "Fulano" in _texto_header(doc)
    assert _texto_footer(doc).strip() == ""


def test_com_logo_valida_embute_imagem_no_cabecalho():
    conteudo = gerar_docx_timbrado(_remetente(logo_base64=_LOGO_PNG))
    doc = Document(io.BytesIO(conteudo))
    assert _tem_imagem_no_header(doc)


def test_com_logo_svg_nao_quebra_e_so_fica_sem_imagem():
    """Mesma limitação da proposta em PDF: fpdf2/python-docx não renderizam
    SVG arbitrário -- sem imagem no cabeçalho, mas nome/documento continuam."""
    svg = "data:image/svg+xml;base64," + "PHN2Zy8+"   # "<svg/>" em base64
    conteudo = gerar_docx_timbrado(_remetente(logo_base64=svg))
    doc = Document(io.BytesIO(conteudo))
    assert not _tem_imagem_no_header(doc)
    assert "Empresa Exemplo LTDA" in _texto_header(doc)


def test_com_logo_valida_cria_marca_dagua_flutuante_atras_do_texto():
    """A marca d'água precisa ser wp:anchor (floating) com behindDoc="1" --
    diferente da logo pequena do cabeçalho, que fica wp:inline (no fluxo do
    texto). Sem isso ela empurraria o layout do cabeçalho em vez de ficar
    atrás, de fundo."""
    conteudo = gerar_docx_timbrado(_remetente(logo_base64=_LOGO_PNG))
    doc = Document(io.BytesIO(conteudo))
    header_xml = doc.sections[0].header._element
    anchors = header_xml.findall(".//" + qn("wp:anchor"))
    assert len(anchors) == 1
    assert anchors[0].get("behindDoc") == "1"


def test_marca_dagua_fica_clara_via_washout():
    """Efeito 'lavado' (luminância alta) na imagem da marca d'água -- sem
    isso ela sairia opaca/escura em cima do texto, ilegível."""
    conteudo = gerar_docx_timbrado(_remetente(logo_base64=_LOGO_PNG))
    doc = Document(io.BytesIO(conteudo))
    header_xml = doc.sections[0].header._element
    luns = header_xml.findall(".//" + qn("a:lum"))
    assert len(luns) == 1
    assert int(luns[0].get("bright")) > 0


def test_sem_logo_nao_gera_marca_dagua():
    conteudo = gerar_docx_timbrado(_remetente(logo_base64=None))
    doc = Document(io.BytesIO(conteudo))
    header_xml = doc.sections[0].header._element
    assert header_xml.findall(".//" + qn("wp:anchor")) == []

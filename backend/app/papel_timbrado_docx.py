"""
Papel timbrado em .docx — mesma identidade visual da proposta em PDF
(proposta_pdf.py: logo + nome no cabeçalho, contato no rodapé, repetidos em
toda página), mas de corpo vazio e editável no Word, pra servir de base
quando o usuário mesmo redige uma declaração (assinatura reconhecida,
ausência de fato impeditivo, etc.) — texto que varia demais de edital pra
edital pra vir pronto do sistema.
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .proposta_pdf import _decodificar_logo

_COR_ACCENT = RGBColor(37, 99, 235)
_COR_MUTED = RGBColor(91, 103, 112)


def _sem_bordas(tabela):
    """python-docx não tem atalho pra remover borda de tabela -- INSERT
    direto no XML (<w:tblBorders> com w:val="nil" em cada lado) é o jeito
    documentado de fazer isso."""
    tbl_pr = tabela._tbl.tblPr
    bordas = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = bordas.makeelement(qn(f"w:{lado}"), {qn("w:val"): "nil"})
        bordas.append(el)
    tbl_pr.append(bordas)


def _marca_dagua(header, buf):
    """Logo grande, clara (efeito 'lavado', tipo o washout do próprio menu
    de marca d'água do Word) e ATRÁS do texto, repetida em toda página --
    mesmo efeito da proposta em PDF (ver proposta_pdf._marca_dagua, que usa
    fill_opacity=0.18). python-docx só sabe criar imagem INLINE (no fluxo do
    texto); marca d'água de verdade precisa ser floating (wp:anchor,
    behindDoc="1") -- não tem atalho na biblioteca, monta-se a mão a partir
    do <wp:inline> que add_picture já gera. Silenciosa em qualquer falha:
    é só estética, nunca pode quebrar a geração do documento."""
    try:
        p = header.add_paragraph()
        run = p.add_run()
        run.add_picture(buf, width=Inches(4.2))
        drawing = run._r.find(qn("w:drawing"))
        inline = drawing.find(qn("wp:inline"))
        extent = inline.find(qn("wp:extent"))
        docpr = inline.find(qn("wp:docPr"))
        cnv = inline.find(qn("wp:cNvGraphicFramePr"))
        graphic = inline.find(qn("a:graphic"))

        anchor = OxmlElement("wp:anchor", {
            "distT": "0", "distB": "0", "distL": "0", "distR": "0",
            "simplePos": "0", "relativeHeight": "1",
            "behindDoc": "1", "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
        })
        simple_pos = OxmlElement("wp:simplePos", {"x": "0", "y": "0"})
        anchor.append(simple_pos)

        pos_h = OxmlElement("wp:positionH", {"relativeFrom": "page"})
        align_h = OxmlElement("wp:align")
        align_h.text = "center"
        pos_h.append(align_h)
        anchor.append(pos_h)

        pos_v = OxmlElement("wp:positionV", {"relativeFrom": "page"})
        align_v = OxmlElement("wp:align")
        align_v.text = "center"
        pos_v.append(align_v)
        anchor.append(pos_v)

        anchor.append(extent)
        anchor.append(OxmlElement("wp:wrapNone"))
        anchor.append(docpr)
        if cnv is not None:
            anchor.append(cnv)
        anchor.append(graphic)
        drawing.replace(inline, anchor)

        # "washout": clareia bastante a imagem -- efeito visual equivalente
        # ao fill_opacity baixo usado na marca d'água do PDF (fpdf2 tem
        # opacidade real; DrawingML não tem esse conceito pra imagem raster,
        # o padrão do próprio Word pra isso é luminância alta + contraste
        # baixo, os mesmos valores que "Inserir > Marca d'água > Washout" usa).
        blip = graphic.find(qn("a:graphicData")).find(qn("pic:pic")) \
                      .find(qn("pic:blipFill")).find(qn("a:blip"))
        blip.append(OxmlElement("a:lum", {"bright": "70000", "contrast": "-70000"}))
    except Exception:
        pass


def gerar_docx_timbrado(remetente: dict) -> bytes:
    """remetente: {nome, documento, endereco: {...}, empresa: {...}, logo_base64}
    -- mesmo formato de _dados_remetente() em main.py."""
    doc = Document()
    secao = doc.sections[0]
    secao.left_margin = secao.right_margin = Inches(0.9)

    # ---------------------------------------------------------------- header
    header = secao.header
    header.is_linked_to_previous = False
    p_cabec = header.paragraphs[0]
    p_cabec.text = ""

    logo = _decodificar_logo(remetente.get("logo_base64"))
    if logo:
        # decodifica de novo (buffer próprio) -- o BytesIO da logo pequena
        # do cabeçalho é consumido pelo add_picture() dela, não dá pra
        # reusar o mesmo objeto pra uma segunda imagem.
        logo_marca = _decodificar_logo(remetente.get("logo_base64"))
        if logo_marca:
            _marca_dagua(header, logo_marca[0])

    tabela = header.add_table(rows=1, cols=2, width=Inches(6.2))
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
    _sem_bordas(tabela)
    cel_logo, cel_texto = tabela.rows[0].cells
    cel_logo.width = Inches(0.9)
    cel_texto.width = Inches(5.3)

    if logo:
        buf, _ext = logo
        try:
            cel_logo.paragraphs[0].add_run().add_picture(buf, width=Inches(0.7))
        except Exception:
            pass

    p_nome = cel_texto.paragraphs[0]
    r_nome = p_nome.add_run(remetente.get("nome") or "")
    r_nome.bold = True
    r_nome.font.size = Pt(13)
    r_nome.font.color.rgb = _COR_ACCENT

    if remetente.get("documento"):
        p_doc = cel_texto.add_paragraph()
        r_doc = p_doc.add_run(f"CNPJ/CPF: {remetente['documento']}")
        r_doc.font.size = Pt(9)
        r_doc.font.color.rgb = _COR_MUTED

    # linha separadora logo abaixo da tabela (parágrafo com borda inferior)
    p_linha = header.add_paragraph()
    p_linha.paragraph_format.space_before = Pt(4)
    p_borda = p_linha._p.get_or_add_pPr()
    bordas_p = p_borda.makeelement(qn("w:pBdr"), {})
    borda_baixo = bordas_p.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "D8DEE6"})
    bordas_p.append(borda_baixo)
    p_borda.append(bordas_p)

    # ---------------------------------------------------------------- footer
    emp = remetente.get("empresa") or {}
    end = remetente.get("endereco") or {}
    partes = [p for p in [
        f"Tel: {emp['telefone']}" if emp.get("telefone") else None,
        ", ".join(v for v in [end.get("logradouro"), end.get("numero"), end.get("bairro")] if v) or None,
        " - ".join(v for v in [end.get("cidade"), end.get("uf")] if v) or None,
        end.get("cep"),
    ] if p]
    if partes:
        p_rodape = secao.footer.paragraphs[0]
        p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_rodape = p_rodape.add_run("  |  ".join(partes))
        r_rodape.font.size = Pt(8)
        r_rodape.font.color.rgb = _COR_MUTED

    # ---------------------------------------------------------- corpo (vazio)
    # Só um parágrafo em branco pra deixar o cursor pronto logo abaixo do
    # cabeçalho -- o texto da declaração em si é redigido pelo usuário.
    doc.add_paragraph()

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
